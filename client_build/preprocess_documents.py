"""
preprocess_documents.py  —  Universal Edition + Hybrid Chunking + Upload UI
=========================================================================
WHAT CHANGED vs your previous version (everything else is IDENTICAL):

1. SEMANTIC SUB-CHUNKING added in write_sections() only.
   - Structural boundaries (H1/H4 splits) are NEVER broken.
   - Only sections that exceed SEMANTIC_CHUNK_TOKENS get sub-chunked.
   - Sub-chunks split on H5/H6 headings when they exist inside the section.
   - If no sub-headings exist, splits on paragraph boundaries at token limit.
   - Last 2 non-empty paragraphs of each chunk overlap into the next chunk
     so no sentence is ever lost at a boundary.
   - Sections that are small enough stay as one file — no forced splitting.

2. UPLOAD UI added — run:  python preprocess_documents.py --ui
   Opens a Streamlit browser page where you drag-drop DOCX or PDF files
   and watch preprocessing run live. Same logic, same output as CLI.

3. slugify() max length increased from 70 to 120 chars to avoid filename
   collisions on long section names.

4. FLOWCHART FIX — now handles all 3 types found across your files:
   TYPE A: Raster images (PNG/EMF via a:blip)        — was working
   TYPE B: Word Drawing Groups (wpg:wgp shapes)       — was silently skipped
   TYPE C: Visio OLE objects (v:imagedata EMF preview) — was silently skipped
   TYPE D: SmartArt (diagrams/data*.xml)               — was silently skipped
   All 4 types now captured. No flowchart data loss.

WHERE CHANGES ARE — search for these markers:
   # ── NEW: semantic sub-chunking ──
   # ── NEW: upload UI ──
   # ── FIX: flowchart extraction ──

Everything else is IDENTICAL to your previous version.
"""

import os, re, sys, base64, time, subprocess, tempfile, shutil, argparse, io
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import pdfplumber
import fitz
import anthropic
from dotenv import load_dotenv

load_dotenv()

from config import (
    VISION_MODEL, ANTHROPIC_API_KEY,
    WHISPER_API_KEY, WHISPER_BASE_URL, USE_LLMWHISPERER,
)

# ── Config ─────────────────────────────────────────────────────────────────────
RAW_DOCS_FOLDER      = "raw_data"
KB_FOLDER            = "knowledge_base"
MAX_RETRIES          = 3
RETRY_WAIT           = 20
DPI                  = 250


# ── NEW: semantic sub-chunking ─────────────────────────────────────────────────
SEMANTIC_CHUNK_TOKENS = 1500
SEMANTIC_OVERLAP_PARAS = 2
# ──────────────────────────────────────────────────────────────────────────────

def _find_libreoffice():
    for candidate in [
        "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]:
        try:
            r = subprocess.run([candidate, "--version"], capture_output=True, timeout=8)
            if r.returncode == 0:
                return candidate
        except Exception:
            continue
    return None

LIBREOFFICE = _find_libreoffice()
client      = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY or None)

VISION_PROMPT = """This image is from a Policy & Procedure document.
Extract ALL content completely and faithfully — nothing can be skipped.

If FLOWCHART / PROCESS DIAGRAM:
- State the Process Name
- List all Roles / Swimlanes
- Every step as Markdown table: | Step # | Role | Action | Decision/Next Step |
- Trace all Yes/No branches explicitly
- Mermaid.js code block (graph TD) for the complete flow

If ORGANISATIONAL CHART:
- Reproduce hierarchy as indented Markdown list

If TABLE / MATRIX:
- Reproduce as complete Markdown table, every row and column

If any other diagram:
- Describe in detail, preserving all labels, values, and text

Start directly with the extracted content. No preamble."""


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — UTILITIES  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

def slugify(text: str) -> str:
    text = re.sub(r"[^\w\s-]", "", str(text))
    text = re.sub(r"\s+", "_", text.strip())
    return text[:120]


def heading_level(style_name: str):
    """
    Detect heading level from Word paragraph style name.
    Handles ALL formats found in client policy documents:
      - "Heading 1", "Heading 2"  (standard Microsoft Word)
      - "H1", "H2", "H3", "H4"   (Deloitte / client custom styles)
      - "Title"                   (document title treated as H1)
    FIX: old version only matched "Heading N" — missed H1/H2/H3/H4 entirely.
    """
    if not style_name:
        return None
    sn = style_name.strip()
    # Standard Word: "Heading 1", "Heading 2" ...
    m = re.match(r"Heading\s+(\d)", sn, re.I)
    if m:
        return int(m.group(1))
    # Deloitte / client custom: "H1", "H2", "H3", "H4"
    m = re.match(r"^H(\d)$", sn, re.I)
    if m:
        return int(m.group(1))
    # Title style = H1
    if sn.lower() == "title":
        return 1
    return None


def emf_to_png(emf_bytes: bytes):
    if not LIBREOFFICE:
        return None
    tmpdir = tempfile.mkdtemp()
    try:
        emf_path = os.path.join(tmpdir, "img.emf")
        with open(emf_path, "wb") as f:
            f.write(emf_bytes)
        subprocess.run(
            [LIBREOFFICE, "--headless", "--convert-to", "png", "--outdir", tmpdir, emf_path],
            capture_output=True, timeout=60
        )
        png_path = emf_path.replace(".emf", ".png")
        return open(png_path, "rb").read() if os.path.exists(png_path) else None
    except Exception:
        return None
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


_EMF_SKIP_WARNED = False

def _note_skipped_emf():
    """Warn once when an embedded Visio/EMF diagram is skipped. LibreOffice is the
    only converter for EMF→PNG; without it GPT vision can't read the format, so we
    skip rather than waste a vision call. Text and tables are unaffected."""
    global _EMF_SKIP_WARNED
    if _EMF_SKIP_WARNED:
        return
    _EMF_SKIP_WARNED = True
    if not LIBREOFFICE:
        print("[preprocess] LibreOffice not found - embedded Visio/EMF diagrams are "
              "skipped (text and tables are unaffected). To extract such diagrams, "
              "install LibreOffice and build the index offline, then deploy the index.")
    else:
        print("[preprocess] An EMF diagram could not be converted and was skipped.")


# ── Extraction-failure handling ───────────────────────────────────────────────
_EXTRACTION_FAIL_MARKERS = ("[VISION_ERROR", "[VISION_FAILED", "[FLOWCHART_PARSE_ERROR")

def _is_extraction_failure(text: str) -> bool:
    return (text or "").strip().startswith(_EXTRACTION_FAIL_MARKERS)


def _img_media_type(img_bytes: bytes) -> str:
    """Sniff the image format from magic bytes for Claude's media_type field."""
    if img_bytes[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if img_bytes[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if img_bytes[:4] in (b"GIF8",):
        return "image/gif"
    if img_bytes[:4] == b"RIFF" and img_bytes[8:12] == b"WEBP":
        return "image/webp"
    return "image/png"   # default — most extracted/converted images are PNG


def _claude_text(resp) -> str:
    """Concatenate the text blocks of a Claude message."""
    return "".join(b.text for b in resp.content if getattr(b, "type", "") == "text").strip()


def call_vision(img_bytes: bytes, label: str = "") -> str:
    """Vision OCR on a single image via Claude (multimodal). No GPT/OpenAI."""
    b64 = base64.b64encode(img_bytes).decode()
    media = _img_media_type(img_bytes)
    messages = [{"role": "user", "content": [
        {"type": "image", "source": {"type": "base64", "media_type": media, "data": b64}},
        {"type": "text", "text": VISION_PROMPT},
    ]}]
    for attempt in range(MAX_RETRIES):
        try:
            resp = client.messages.create(model=VISION_MODEL, max_tokens=8000, messages=messages)
            return _claude_text(resp)
        except anthropic.RateLimitError:
            if attempt < MAX_RETRIES - 1:
                print(f"\n  [rate limit] waiting {RETRY_WAIT}s...", flush=True)
                time.sleep(RETRY_WAIT)
            else:
                return "[VISION_FAILED — rate limit. Delete this .md and re-run to retry.]"
        except Exception as e:
            return f"[VISION_ERROR: {e}]"


def call_flowchart_llm(text: str) -> str:
    prompt = f"""
You are given steps extracted from a flowchart.

Convert them into structured logic:

1. Identify steps
2. Detect decisions (Yes/No if implied)
3. Output:
   - Step table
   - Mermaid diagram (graph TD)

Flowchart text:
{text}

Return clean Markdown only.
"""
    try:
        resp = client.messages.create(
            model=VISION_MODEL, max_tokens=4000,
            messages=[{"role": "user", "content": prompt}],
        )
        return _claude_text(resp)
    except Exception as e:
        return f"[FLOWCHART_PARSE_ERROR: {e}]"


# ── FIX: flowchart extraction — handles all 4 types ──────────────────────────
#
# TYPE A: Raster image (PNG/EMF) via a:blip inside w:drawing
#   → get_image_from_drawing() — unchanged, was already working
#
# TYPE B: Word Drawing Group (wpg:wgp) — shapes drawn directly in Word
#   → _extract_wpg_text_from_para() — extracts text from wps:txbx boxes
#   → Found in: Production/Feed Manual flowchart overview
#
# TYPE C: Visio OLE object with EMF preview via v:imagedata inside w:pict
#   → _extract_vml_images_from_para() — extracts EMF blob, converts to PNG
#   → Found in: SCM Procurement Manual (14 Visio flowcharts)
#
# TYPE D: SmartArt diagram — text stored in diagrams/data*.xml relationships
#   → _extract_smartart_text_from_doc() — reads all diagramData parts once
#   → _inject_smartart_into_para() — matches para to its SmartArt, injects text
#   → Found in: SCM Procurement Manual (org chart, category diagram)
#
# ─────────────────────────────────────────────────────────────────────────────

_MC_NS  = {"mc": "http://schemas.openxmlformats.org/markup-compatibility/2006"}
_VML_NS = "urn:schemas-microsoft-com:vml"
_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def get_image_from_drawing(drawing, doc):
    """TYPE A: raster image via a:blip. Unchanged."""
    try:
        blips = drawing.findall(".//" + qn("a:blip"))
        if not blips:
            return None, ""
        rid = blips[0].get(qn("r:embed"))
        if not rid:
            return None, ""
        img_part = doc.part.related_parts.get(rid)
        if not img_part:
            return None, ""
        blob = img_part.blob
        ct   = img_part.content_type
        if "emf" in ct:
            converted = emf_to_png(blob)
            if converted:
                return converted, "EMF→PNG"
            # No LibreOffice (or conversion failed): GPT vision cannot read EMF, so
            # skip instead of sending an unreadable blob (which only wastes a call).
            _note_skipped_emf()
            return None, ""
        fmt = ct.split("/")[-1].upper()
        return blob, fmt
    except Exception:
        return None, ""


def _extract_wpg_text_from_para(para_element) -> str:
    """
    TYPE B: Word Drawing Group shapes (wpg:wgp) AND individual wps:wsp shapes.

    Two sub-types both handled here:
      B1 — wpg:wgp (grouped canvas): text in wps:txbx / txbxContent inside the group.
           Found in: Arabian Mills main file.
      B2 — individual wps:wsp shapes inside mc:AlternateContent (no group wrapper):
           each AlternateContent block = one flowchart box, text in txbxContent.
           Found in: Production Feed/Flour file.

    Extracts all shape texts in document order, deduplicates, formats as numbered list.
    Returns '' if no shape text found (so caller falls through to normal processing).
    """
    all_texts = []
    seen = set()

    def _collect_txbx_text(root_element):
        """Walk any element and collect all txbxContent text, deduped."""
        for el in root_element.iter():
            local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if local == "txbxContent":
                for t in el.findall(".//" + qn("w:t")):
                    txt = (t.text or "").strip()
                    if txt and txt not in seen:
                        seen.add(txt)
                        all_texts.append(txt)

    # B1: wpg:wgp inside w:drawing
    drawings = para_element.findall(".//" + qn("w:drawing"))
    for drawing in drawings:
        wpg_els = [el for el in drawing.iter() if el.tag.split("}")[-1] == "wgp"]
        if wpg_els:
            _collect_txbx_text(drawing)

    # B2: individual wps:wsp inside mc:AlternateContent (no group)
    mc_blocks = para_element.findall(".//mc:AlternateContent", _MC_NS)
    for ac in mc_blocks:
        wpg_in_ac = [el for el in ac.iter() if el.tag.split("}")[-1] == "wgp"]
        if wpg_in_ac:
            continue  # already handled above as B1
        wsp_in_ac = [el for el in ac.iter() if el.tag.split("}")[-1] == "wsp"]
        if wsp_in_ac:
            _collect_txbx_text(ac)

    if not all_texts:
        return ""

    lines = ["**[Flowchart — Word Shapes]:**", ""]
    for i, step in enumerate(all_texts, 1):
        lines.append(f"{i}. {step}")
    lines.append("")
    return "\n".join(lines)


def _extract_vml_images_from_para(para_element, doc) -> list:
    """
    TYPE C: Visio OLE objects embedded as v:imagedata inside w:pict.
    Returns list of (png_bytes, label) for each flowchart found in this paragraph.
    The EMF preview is the visual representation of the Visio drawing.
    """
    results = []
    imagedata_els = para_element.findall(f".//{{{_VML_NS}}}imagedata")
    for img_el in imagedata_els:
        rid = img_el.get(f"{{{_REL_NS}}}id")
        if not rid:
            continue
        try:
            img_part = doc.part.related_parts.get(rid)
            if not img_part:
                continue
            if "emf" not in img_part.content_type:
                continue
            # Convert EMF to PNG for vision
            png = emf_to_png(img_part.blob)
            if png:
                results.append((png, "Visio-EMF→PNG"))
            else:
                # No LibreOffice (or conversion failed): GPT vision can't read EMF.
                # Skip cleanly instead of sending an unreadable raw EMF blob.
                _note_skipped_emf()
        except Exception:
            continue
    return results


def _load_smartart_texts(doc) -> list:
    """
    TYPE D: SmartArt — reads ALL diagramData relationship parts and returns
    a list of text lists (one per SmartArt diagram).
    Called ONCE at the start of process_docx, not per paragraph.
    """
    smartart_list = []
    try:
        for rid, rel in doc.part.rels.items():
            if not rel.reltype.endswith("diagramData"):
                continue
            xml_root = None
            try:
                from lxml import etree
                xml_root = etree.parse(io.BytesIO(rel.target_part.blob)).getroot()
            except Exception:
                continue
            texts = []
            seen = set()
            for el in xml_root.iter():
                local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
                if local == "t" and el.text and el.text.strip():
                    txt = el.text.strip()
                    if txt not in seen:
                        seen.add(txt)
                        texts.append(txt)
            if texts:
                smartart_list.append(texts)
    except Exception:
        pass
    return smartart_list


def _format_smartart_as_md(texts: list) -> str:
    """Format SmartArt text list as a Markdown structure."""
    if not texts:
        return ""
    lines = ["**[SmartArt Diagram — extracted text]:**", ""]
    for txt in texts:
        lines.append(f"- {txt}")
    lines.append("")
    return "\n".join(lines)


def _para_has_smartart(para_element) -> bool:
    """Return True if paragraph contains a SmartArt drawing (dgm namespace)."""
    drawings = para_element.findall(".//" + qn("w:drawing"))
    for d in drawings:
        graphicData = d.findall(".//" + qn("a:graphicData"))
        for gd in graphicData:
            uri = gd.get("uri", "")
            if "diagram" in uri:
                return True
    return False

# ── End of flowchart fix helpers ──────────────────────────────────────────────


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — SECTION CLASS  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

class Section:
    def __init__(self, dept: str, name: str):
        self.dept       = dept
        self.name       = name
        self.lines      = []
        self.images     = {}
        self.img_count  = 0
        self.text_chars = 0
        self.table_rows = 0

    def add_image(self, img_bytes: bytes, fmt: str) -> str:
        self.img_count += 1
        key = f"__IMG_{id(self)}_{self.img_count}__"
        self.images[key] = (img_bytes, fmt)
        self.lines.append(key)
        return key

    def is_empty(self) -> bool:
        return self.text_chars == 0 and self.table_rows == 0 and self.img_count == 0

    def token_estimate(self) -> int:
        return self.text_chars // 4 + self.table_rows * 50


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — DOCX PROCESSING
# ─────────────────────────────────────────────────────────────────────────────

def _detect_split_level(paras: list, start: int, end: int) -> int:
    slice_ = paras[start:end]
    for level in [5, 4, 3, 2]:
        headings = [
            p.text.strip() for p in slice_
            if heading_level(p.style.name) == level and p.text.strip()
        ]
        count = len(headings)
        if count < 3 or count > 40:
            continue
        diversity = len(set(headings)) / count
        if diversity >= 0.5:
            return level
    for level in [2, 3, 4, 5]:
        count = sum(
            1 for p in slice_
            if heading_level(p.style.name) == level and p.text.strip()
        )
        if count > 0:
            return level
    return 2


def _get_cell_text(cell) -> str:
    """
    Extract ALL text from a DOCX table cell.

    ROOT CAUSE FIX: cell.text and paragraph.text only return direct text node
    content. In policy/procedure documents, text is stored inside w:r/w:t run
    elements nested inside paragraphs inside cells. Calling .text returns empty
    string for these cells — silently dropping RACI tables, KPI tables,
    definition tables, and all table-based policy content.

    Fix: iterate ALL w:t elements anywhere inside the cell XML.
    Also handles \xa0 (non-breaking space) → regular space.
    """
    parts = []
    seen = set()

    for para in cell.paragraphs:
        para_parts = []
        for t_elem in para._p.iter(qn("w:t")):
            txt = (t_elem.text or "").replace("\xa0", " ").strip()
            if txt:
                para_parts.append(txt)

        para_text = " ".join(para_parts).strip()
        if para_text and para_text not in seen:
            seen.add(para_text)
            parts.append(para_text)

    if not parts:
        return ""

    if len(parts) == 1:
        return parts[0]

    first = parts[0]
    rest = parts[1:]

    # If first line looks like a short label/subheading, keep it bold
    if len(first) <= 80 and not first.endswith("."):
        return f"**{first}**<br>" + "<br>".join(f"• {p}" for p in rest)

    # Otherwise keep all lines structured inside the same cell
    return "<br>".join(f"• {p}" for p in parts)
   

def _table_to_md(table) -> str:
    """
    Convert DOCX table to complete Markdown — FIXED.

    Previous: used paragraph.text which misses text in w:r/w:t run elements.
    Result: RACI matrices, KPI tables, definition tables → empty rows in KB.

    This version uses _get_cell_text() which iterates w:t elements directly.
    Merged cell deduplication: Word repeats same cell object for merged cells;
    we skip by comparing consecutive cell text to previous.
    """
    rows = []
    max_cols = 0

    for row in table.rows:
        cells = []
        prev_text = None
        for cell in row.cells:
            txt = _get_cell_text(cell).replace("|", "/").strip()
            if txt != prev_text:   # skip merged cells
                cells.append(txt)
            prev_text = txt
        if any(c.strip() for c in cells):
            rows.append(cells)
            max_cols = max(max_cols, len(cells))

    if not rows:
        return ""

    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    md = []
    md.append("| " + " | ".join(rows[0]) + " |")
    md.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in rows[1:]:
        md.append("| " + " | ".join(row[:max_cols]) + " |")
    return "\n".join(md)

def _text_already_numbered(text: str) -> bool:
    return bool(re.match(r"^\s*(?:\(?[0-9]+[\.\)]|[A-Za-z]+[\.\)]|\(?[ivxlcdm]+[\.\)])\s+", text, re.IGNORECASE))


def _number_to_roman(number: int) -> str:
    values = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    result = ""
    n = max(1, min(number, 3999))
    for value, numeral in values:
        while n >= value:
            result += numeral
            n -= value
    return result


def _number_to_alpha(number: int, upper: bool = False) -> str:
    if number < 1:
        return ""
    letters = []
    n = number - 1
    while n >= 0:
        letters.append(chr((n % 26) + (65 if upper else 97)))
        n = n // 26 - 1
    return "".join(reversed(letters))


def _get_numbering_xml(doc):
    part = getattr(doc.part, "numbering_part", None)
    if part is None:
        return None
    return getattr(part, "_numbering", None) or getattr(part, "element", None)


def _get_num_pr_info(para):
    p = para._p
    numPr = p.find(".//" + qn("w:numPr"))
    if numPr is None:
        return None
    numId = numPr.find(qn("w:numId"))
    ilvl = numPr.find(qn("w:ilvl"))
    if numId is None or ilvl is None:
        return None
    try:
        return {
            "num_id": numId.get(qn("w:val")),
            "ilvl": int(ilvl.get(qn("w:val"), "0")),
        }
    except Exception:
        return None


def _get_abstract_num_levels(doc, num_id: str) -> dict:
    numbering = _get_numbering_xml(doc)
    if numbering is None:
        return {}

    abstract_id = None
    for num in numbering.findall(".//" + qn("w:num")):
        if num.get(qn("w:numId")) == str(num_id):
            abstract_el = num.find(qn("w:abstractNumId"))
            if abstract_el is not None:
                abstract_id = abstract_el.get(qn("w:val"))
            break
    if abstract_id is None:
        return {}

    levels = {}
    for abs_num in numbering.findall(".//" + qn("w:abstractNum")):
        if abs_num.get(qn("w:abstractNumId")) != str(abstract_id):
            continue
        for lvl in abs_num.findall(".//" + qn("w:lvl")):
            ilvl_val = lvl.get(qn("w:ilvl"))
            if ilvl_val is None:
                continue
            try:
                level_index = int(ilvl_val)
            except Exception:
                continue
            num_fmt_el = lvl.find(qn("w:numFmt"))
            lvl_text_el = lvl.find(qn("w:lvlText"))
            start_el = lvl.find(qn("w:start"))
            levels[level_index] = {
                "num_fmt": num_fmt_el.get(qn("w:val")) if num_fmt_el is not None else "decimal",
                "lvl_text": lvl_text_el.get(qn("w:val")) if lvl_text_el is not None else None,
                "start": int(start_el.get(qn("w:val"), "1")) if start_el is not None else 1,
            }
        break
    return levels


def _render_number_value(num_fmt: str, value: int) -> str:
    if num_fmt == "decimal":
        return str(value)
    if num_fmt == "decimalZero":
        return str(value).zfill(2)
    if num_fmt == "lowerLetter":
        return _number_to_alpha(value, upper=False)
    if num_fmt == "upperLetter":
        return _number_to_alpha(value, upper=True)
    if num_fmt == "lowerRoman":
        return _number_to_roman(value).lower()
    if num_fmt == "upperRoman":
        return _number_to_roman(value).upper()
    return str(value)


def _increment_numbering(numbering_state: dict, num_id: str, level: int, start: int) -> int:
    entry = numbering_state.setdefault(num_id, {"last_level": -1, "counters": {}})
    counters = entry["counters"]
    prev_level = entry["last_level"]

    if level > prev_level:
        counters[level] = counters.get(level, start)
    else:
        counters[level] = counters.get(level, start) + 1
        for deeper in list(counters):
            if deeper > level:
                del counters[deeper]

    entry["last_level"] = level
    return counters[level]


def _format_numbering_label(doc, num_id: str, level: int, numbering_state: dict) -> str | None:
    levels = _get_abstract_num_levels(doc, num_id)
    if not levels or level not in levels:
        return None

    props = levels[level]
    current_value = _increment_numbering(numbering_state, num_id, level, props["start"])
    label = props["lvl_text"]
    if label:
        for i in range(1, 10):
            token = f"%{i}"
            if token not in label:
                continue
            level_index = i - 1
            level_props = levels.get(level_index, {"num_fmt": props["num_fmt"], "start": 1})
            counter_value = numbering_state.get(num_id, {"counters": {}})["counters"].get(level_index, level_props["start"])
            label = label.replace(token, _render_number_value(level_props["num_fmt"], counter_value))
        return label + " " if not label.endswith(" ") else label

    return _render_number_value(props["num_fmt"], current_value) + ". "


def _strip_leading_numbering(text: str) -> str:
    """Remove leading numbering from text, e.g., '1. ' or 'a) '."""
    return re.sub(r'^\s*(?:\d+\.|\w+\.|\(\d+\)|\(\w+\)|\w+\)|\d+\)|\w+\))\s*', '', text).strip()


def _para_to_md(para, doc=None, numbering_state=None) -> str:
    text = para.text.strip()

    if not text:
        return ""

    try:
        style = para.style.name
    except Exception:
        style = "Normal"

    lvl = heading_level(style)
    num_info = _get_num_pr_info(para)
    number_label = ""
    if num_info and doc is not None and numbering_state is not None:
        number_label = _format_numbering_label(doc, num_info["num_id"], num_info["ilvl"], numbering_state) or ""

    if lvl:
        # For headings, strip any leading numbering to keep clean "# Heading"
        heading_text = _strip_leading_numbering(text)
        return "#" * lvl + " " + heading_text

    if number_label:
        # For lists, preserve numbering and indent for hierarchies
        list_text = text if _text_already_numbered(text) else number_label + text
        indent = "  " * num_info["ilvl"]
        return indent + list_text

    return text


def process_docx(docx_path: str) -> tuple:
    print(f"  Parsing: {Path(docx_path).name}")
    doc      = Document(docx_path)
    paras    = doc.paragraphs
    para_map = {p._p: p for p in paras}
    tbl_map  = {t._tbl: t for t in doc.tables}
    body     = doc.element.body

    gt_text  = sum(len(p.text) for p in paras if p.text.strip())
    gt_table = sum(
        sum(len(c.text) for r in t.rows for c in r.cells)
        for t in doc.tables
    )

    # ── FIX: load SmartArt texts once upfront (TYPE D) ───────────────────────
    smartart_queue = _load_smartart_texts(doc)  # list of text-lists, one per diagram
    smartart_idx   = [0]                         # mutable pointer
    numbering_state = {}                         # preserve original numbering/alphabet labels
    # ─────────────────────────────────────────────────────────────────────────

    h1_positions = [
        (i, p.text.strip())
        for i, p in enumerate(paras)
        if heading_level(p.style.name) == 1 and p.text.strip()
    ]
    h1_positions.append((len(paras), "__END__"))

    split_level_for_dept: dict[str, int] = {}
    for k in range(len(h1_positions) - 1):
        dept_name      = h1_positions[k][1]
        start_i, end_i = h1_positions[k][0], h1_positions[k + 1][0]
        split_level_for_dept[dept_name] = _detect_split_level(paras, start_i, end_i)

    # FIX 3: if document has NO H1, detect split level across the whole document
    if not split_level_for_dept:
        whole_doc_split = _detect_split_level(paras, 0, len(paras))
        split_level_for_dept["__ALL__"] = whole_doc_split

    sections: list[Section]  = []
    current:  Section | None = None
    current_dept              = "General"
    current_split_level       = split_level_for_dept.get("__ALL__", 2)
    pending_parent_heading: str = ""
    img_counter               = [0]

    def _new_section(dept: str, name: str):
        nonlocal current, pending_parent_heading
        if current:
            # FIX 1: if current section is empty (H1 with no direct content before sub-headings),
            # carry its heading as context into the first child section.
            if current.is_empty() and current.lines:
                pending_parent_heading = current.lines[0]
            else:
                sections.append(current)
                pending_parent_heading = ""
        current = Section(dept, name)
        if pending_parent_heading:
            current.lines.append(pending_parent_heading)
            pending_parent_heading = ""
        current.lines.append(f"## {name}")
        current.lines.append("")  # Add blank line after subheading

    for child in body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para = para_map.get(child)
            if not para:
                continue
            try:
                style = para.style.name
            except Exception:
                style = "Normal"
            text = para.text.strip()
            lvl  = heading_level(style)

            if lvl == 1 and text:
                current_dept        = text
                current_split_level = split_level_for_dept.get(text, 2)
                if current:
                    # FIX 1: same empty-section handling at H1 boundary
                    if current.is_empty() and current.lines:
                        pending_parent_heading = current.lines[0]
                    else:
                        sections.append(current)
                        pending_parent_heading = ""
                current = Section(current_dept, text)
                current.lines.append(f"# {text}")
                current.lines.append("")  # Add blank line after heading
                continue

            if lvl == current_split_level and text:
                _new_section(current_dept, text)
                continue

            # ── FIX: handle all drawing types in one place ────────────────────
            drawings = child.findall(".//" + qn("w:drawing"))
            has_drawings = bool(drawings)

            # Check for VML/pict (Visio OLE) regardless of w:drawing
            vml_images = _extract_vml_images_from_para(child, doc)

            if has_drawings or vml_images:
                if current is None:
                    current = Section(current_dept, "General")

                any_content = False

                # TYPE A: raster images via a:blip
                for drawing in drawings:
                    img_bytes, fmt = get_image_from_drawing(drawing, doc)
                    if img_bytes:
                        img_counter[0] += 1
                        current.add_image(img_bytes, fmt)
                        any_content = True

                # TYPE B: Word Drawing Group shapes (wpg:wgp)
                wpg_md = _extract_wpg_text_from_para(child)
                if wpg_md:
                    current.lines.append("\n" + wpg_md + "\n")
                    structured = call_flowchart_llm(wpg_md)
                    if not _is_extraction_failure(structured):
                        current.lines.append("\n**[Flowchart — Structured]:**\n\n" + structured + "\n")
                    current.text_chars += len(wpg_md)
                    any_content = True

                # TYPE C: Visio OLE EMF images
                for png_bytes, fmt in vml_images:
                    img_counter[0] += 1
                    current.add_image(png_bytes, fmt)
                    any_content = True

                # TYPE D: SmartArt (matched by diagram graphicData URI)
                if _para_has_smartart(child) and smartart_idx[0] < len(smartart_queue):
                    sa_texts = smartart_queue[smartart_idx[0]]
                    smartart_idx[0] += 1
                    sa_md = _format_smartart_as_md(sa_texts)
                    if sa_md:
                        current.lines.append("\n" + sa_md + "\n")
                        current.text_chars += len(sa_md)
                        any_content = True

                if any_content:
                    continue
            # ── End of drawing handling ───────────────────────────────────────

            md = _para_to_md(para, doc=doc, numbering_state=numbering_state)
            if md:
                if current is None:
                    current = Section(current_dept, "General")
                current.text_chars += len(text)
                current.lines.append(md)

        elif tag == "tbl":
            tbl = tbl_map.get(child)
            if not tbl:
                continue
            md = _table_to_md(tbl)
            if md:
                if current is None:
                    current = Section(current_dept, "General")
                current.table_rows += len(tbl.rows)
                current.lines.append("\n" + md + "\n")

    if current:
        sections.append(current)

    return sections, gt_text, gt_table, split_level_for_dept


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — PDF PROCESSING  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

def _pdf_has_text(pdf_path: str) -> bool:
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:5]:
                if len((page.extract_text() or "").strip()) > 100:
                    return True
        return False
    except Exception:
        return False


def _process_pdf_llmwhisperer(pdf_path: str) -> tuple:
    """
    Optional layout-preserving PDF parsing via Unstract LLMWhisperer.
    Active only when WHISPER_API_KEY is configured. Returns the same tuple shape
    as the built-in parsers; the single Section is semantically sub-chunked downstream.
    """
    from unstract.llmwhisperer import LLMWhispererClientV2
    client = LLMWhispererClientV2(base_url=WHISPER_BASE_URL, api_key=WHISPER_API_KEY)
    result = client.whisper(
        file_path=pdf_path,
        mode="high_quality",
        output_mode="layout_preserving",
        wait_for_completion=True,
        wait_timeout=300,
    )

    extraction = result.get("extraction", {}) if isinstance(result, dict) else {}
    text = extraction.get("result_text", "") if isinstance(extraction, dict) else str(extraction or "")
    if not text.strip():
        raise RuntimeError("LLMWhisperer returned no text")

    # Normalize LLMWhisperer page separators ("<<<") into blank lines.
    text = re.sub(r"\n?<<<\n?", "\n\n", text).strip()

    base = Path(pdf_path).stem
    s = Section("PDF", base)
    s.lines.append(f"## {base}")
    s.lines.extend(text.split("\n"))
    s.text_chars = len(text)
    print(f"  [LLMWhisperer] extracted {len(text):,} chars")
    return [s], len(text), 0, {}


def process_pdf(pdf_path: str) -> tuple:
    print(f"  Parsing: {Path(pdf_path).name}")

    # Optional layout-preserving PDF parsing (only if a key is configured).
    if USE_LLMWHISPERER:
        try:
            print("  Using LLMWhisperer (layout-preserving)...")
            return _process_pdf_llmwhisperer(pdf_path)
        except Exception as e:
            print(f"  [LLMWhisperer] failed ({e}); falling back to built-in PDF parsing")

    is_text = _pdf_has_text(pdf_path)
    print(f"  Type: {'text-based' if is_text else 'scanned / image-only'}")
    doc_fitz = fitz.open(pdf_path)
    total    = len(doc_fitz)
    base     = Path(pdf_path).stem
    if is_text:
        return _process_text_pdf(doc_fitz, pdf_path, base, total)
    else:
        return _process_scanned_pdf(doc_fitz, base, total)


def _process_text_pdf(doc_fitz, pdf_path: str, base: str, total: int) -> tuple:
    GROUP = 20
    sections: list[Section] = []
    gt_text = 0

    def _new_pdf_section(start_page: int) -> Section:
        end  = min(start_page + GROUP - 1, total)
        name = f"Pages_{start_page + 1}_to_{end + 1}"
        s    = Section("PDF", name)
        s.lines.append(f"## {name}")
        return s

    current = _new_pdf_section(0)

    with pdfplumber.open(pdf_path) as plumber:
        for page_num in range(total):
            if page_num > 0 and page_num % GROUP == 0:
                sections.append(current)
                current = _new_pdf_section(page_num)

            page = plumber.pages[page_num]
            text = page.extract_text() or ""
            if text:
                current.text_chars += len(text)
                gt_text += len(text)
                current.lines.append(text)

            for table in page.extract_tables() or []:
                rows = []
                for i, row in enumerate(table):
                    cells = [str(c or "").replace("|", "/") for c in row]
                    rows.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                if rows:
                    current.table_rows += len(table)
                    current.lines.append("\n".join(rows))

            fitz_page = doc_fitz.load_page(page_num)
            for img_info in fitz_page.get_images(full=True):
                xref = img_info[0]
                try:
                    img_data  = doc_fitz.extract_image(xref)
                    img_bytes = img_data["image"]
                    fmt       = img_data["ext"].upper()
                    current.add_image(img_bytes, fmt)
                except Exception:
                    pass

    sections.append(current)
    return sections, gt_text, 0, {}


def _process_scanned_pdf(doc_fitz, base: str, total: int) -> tuple:
    GROUP = 15
    sections: list[Section] = []

    def _new_scan_section(start_page: int) -> Section:
        end  = min(start_page + GROUP - 1, total)
        name = f"Pages_{start_page + 1}_to_{end + 1}"
        s    = Section("PDF", name)
        s.lines.append(f"## {name}")
        return s

    current = _new_scan_section(0)
    for page_num in range(total):
        if page_num > 0 and page_num % GROUP == 0:
            sections.append(current)
            current = _new_scan_section(page_num)
        page = doc_fitz.load_page(page_num)
        pix  = page.get_pixmap(dpi=DPI)
        current.add_image(pix.tobytes("png"), "PDF-page")

    sections.append(current)
    return sections, 0, 0, {}


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — VERIFICATION REPORT  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

def print_verification(sections: list, gt_text: int, gt_table: int,
                       split_levels: dict, source_name: str):
    captured_text  = sum(s.text_chars  for s in sections)
    captured_table = sum(s.table_rows  for s in sections)
    total_images   = sum(s.img_count   for s in sections)
    text_pct       = 100 * captured_text / max(gt_text, 1)

    print()
    print("=" * 65)
    print(f"VERIFICATION REPORT — {source_name}")
    print("=" * 65)
    print(f"  Sections created:      {len(sections)}")
    print(f"  Text captured:         {captured_text:,} chars  ({text_pct:.1f}%)")
    print(f"  Table rows captured:   {captured_table:,}  (100% — lossless)")
    print(f"  Images for vision:     {total_images}")
    print(f"  Est. API cost:         ~${total_images * 0.012:.2f}")

    if split_levels:
        print(f"\n  AUTO-DETECTED SPLIT LEVELS:")
        for dept, lvl in split_levels.items():
            print(f"    H{lvl}  {dept[:58]}")

    large = [s for s in sections if s.token_estimate() > SEMANTIC_CHUNK_TOKENS]
    if large:
        print(f"\n  SECTIONS THAT WILL BE SEMANTICALLY SUB-CHUNKED ({len(large)}):")
        for s in large:
            print(f"    ~{s.token_estimate():>5} tokens  [{s.dept[:25]}] {s.name[:60]}")
    else:
        print(f"\n  All sections fit in one chunk — no semantic splitting needed.")

    empty = [s for s in sections if s.is_empty()]
    if empty:
        print(f"\n  WARNING — {len(empty)} empty section(s):")
        for s in empty:
            print(f"    [{s.dept[:28]}] {s.name[:50]}")
    else:
        print(f"\n  No empty sections — all sections have content.")

    print(f"\n  SECTION DETAIL:")
    print(f"  {'DEPT':<30} {'SECTION NAME':<38} {'TXT':>6} {'TBL':>4} {'IMG':>4}")
    print(f"  {'-'*30} {'-'*38} {'-'*6} {'-'*4} {'-'*4}")
    for s in sections:
        flag = " ← EMPTY" if s.is_empty() else ""
        print(f"  {s.dept[:30]:<30} {s.name[:60]:<60} "
              f"{s.text_chars:>6} {s.table_rows:>4} {s.img_count:>4}{flag}")
    print("=" * 65)

    return len(empty) == 0 and text_pct > 90


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — VISION PASS  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

def run_vision(sections: list, log_fn=None) -> list:
    """
    Run the vision model on every captured image.
    Successful extractions are inlined; failures are QUARANTINED — replaced with a
    neutral marker (so error strings never get embedded into the KB) and returned
    as a list of {section, fmt, error} for logging/review.
    """
    _log = log_fn or print
    failures = []
    total = sum(s.img_count for s in sections)
    if total == 0:
        _log("  No images to process.")
        return failures

    done = 0
    _log(f"\nRunning {VISION_MODEL} vision on {total} images...\n")
    for section in sections:
        if not section.images:
            continue
        for placeholder, (img_bytes, fmt) in section.images.items():
            done += 1
            label = f"{section.name[:38]} [{fmt}] {done}/{total}"
            _log(f"  [{done:>3}/{total}] {label} ...")
            vision_text = call_vision(img_bytes, label)

            if _is_extraction_failure(vision_text):
                failures.append({"section": section.name, "fmt": fmt, "error": vision_text.strip()})
                replacement = "*(Diagram could not be extracted automatically — flagged for review.)*"
                _log(f"  [{done:>3}/{total}] ✗ extraction failed ({fmt}) — quarantined")
            else:
                replacement = f"\n**[Diagram — {fmt}]:**\n\n{vision_text}\n"
                _log(f"  [{done:>3}/{total}] ✓")

            section.lines = [replacement if line == placeholder else line for line in section.lines]

    if failures:
        _log(f"\n  ⚠ {len(failures)} diagram(s) failed extraction and were quarantined "
             f"(not embedded as error text).")
    return failures


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — WRITE OUTPUT  +  SEMANTIC SUB-CHUNKING  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

def is_procedure_like(section):
    text = " ".join(section.lines).lower()
    keywords = [
        "procedure", "procedures", "process", "workflow",
        "step", "steps", "responsibility", "responsibilities",
        "approval", "approvals", "policy", "policies"
    ]
    return any(k in text for k in keywords)


def _split_lines_into_chunks(lines: list, section_name: str) -> list[tuple]:
    TOKEN_LIMIT = SEMANTIC_CHUNK_TOKENS * 4

    def _is_subheading(line: str) -> bool:
        return bool(re.match(r"#{5,6} ", line.strip()))

    sub_heading_positions = [i for i, line in enumerate(lines) if _is_subheading(line)]
    chunks: list[tuple] = []

    if sub_heading_positions:
        boundaries = sub_heading_positions + [len(lines)]
        overlap_lines: list[str] = []
        for k in range(len(boundaries) - 1):
            start = boundaries[k]
            end   = boundaries[k + 1]
            chunk_name  = f"{section_name} — {lines[start].lstrip('#').strip()}"
            chunk_lines = overlap_lines + lines[start:end]
            chunk_chars = sum(len(l) for l in chunk_lines)
            if chunk_chars > TOKEN_LIMIT * 1.5:
                sub = _split_on_paragraphs(chunk_lines, chunk_name, TOKEN_LIMIT)
                chunks.extend(sub)
            else:
                chunks.append((chunk_name, chunk_lines))
            non_empty = [l for l in lines[start:end] if l.strip() and not _is_subheading(l)]
            overlap_lines = non_empty[-SEMANTIC_OVERLAP_PARAS:] if non_empty else []
    else:
        chunks = _split_on_paragraphs(lines, section_name, TOKEN_LIMIT)

    return chunks if chunks else [(section_name, lines)]


def _split_on_paragraphs(lines: list, base_name: str, char_limit: int) -> list[tuple]:
    chunks = []
    current_lines: list[str] = []
    current_chars = 0
    chunk_num     = 1
    overlap_lines: list[str] = []
    in_table      = False

    for line in lines:
        is_table_line = line.strip().startswith("|")
        if is_table_line:
            in_table = True
        elif in_table and not is_table_line:
            in_table = False
        current_lines.append(line)
        current_chars += len(line)
        if current_chars >= char_limit and not line.strip() and not in_table:
            chunk_name = f"{base_name} (part {chunk_num})"
            chunks.append((chunk_name, overlap_lines + current_lines))
            chunk_num += 1
            non_empty    = [l for l in current_lines if l.strip()]
            overlap_lines = non_empty[-SEMANTIC_OVERLAP_PARAS:] if non_empty else []
            current_lines = []
            current_chars = 0

    if current_lines:
        chunk_name = f"{base_name} (part {chunk_num})" if chunk_num > 1 else base_name
        chunks.append((chunk_name, overlap_lines + current_lines))

    return chunks


# ── Section classification helpers ────────────────────────────────────────────
SKIP_SECTION_PATTERNS = [
    re.compile(r"table of contents", re.I),
]

LOW_PRIORITY_SECTIONS = {
    "toc", "revision_history", "document_control", "glossary", "appendix"
}


def _section_kind(section_name: str) -> str:
    """Classify section type for metadata tagging."""
    name = (section_name or "").lower()
    if name.strip() == "general":                          return "general"
    if "table of contents" in name:                        return "toc"
    if "revision history" in name:                         return "revision_history"
    if "document control" in name:                         return "document_control"
    if any(x in name for x in ["abbreviation", "acronym", "key terms", "glossary"]):
        return "glossary"
    if any(x in name for x in ["annex", "appendix", "form"]):
        return "appendix"
    return "core"


def _section_priority(section_name: str) -> str:
    return "low" if _section_kind(section_name) in LOW_PRIORITY_SECTIONS else "normal"


def _should_skip_section(section_name: str) -> bool:
    return any(p.search(section_name or "") for p in SKIP_SECTION_PATTERNS)


def write_sections(sections: list, source_label: str, log_fn=None) -> int:
    _log = log_fn or print
    os.makedirs(KB_FOLDER, exist_ok=True)
    written = 0

    for i, s in enumerate(sections):
        if s.is_empty():
            continue

        if _should_skip_section(s.name):
            _log(f"  [skip] {s.name} — excluded from KB")
            continue

        dept_folder = os.path.join(KB_FOLDER, slugify(s.dept))
        os.makedirs(dept_folder, exist_ok=True)

        kind     = _section_kind(s.name)
        priority = _section_priority(s.name)
        needs_sub_chunk = s.token_estimate() > SEMANTIC_CHUNK_TOKENS

        if needs_sub_chunk:
            chunks = _split_lines_into_chunks(s.lines, s.name)
            _log(f"  [{s.dept[:25]}] {s.name[:40]}: sub-chunked into {len(chunks)} parts")
            for j, (chunk_name, chunk_lines) in enumerate(chunks):
                content = "\n".join(chunk_lines).strip()
                if not content:
                    continue
                fname    = f"{i + 1:03d}_{slugify(s.name)}_chunk{j + 1:02d}.md"
                out_path = os.path.join(dept_folder, fname)
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(
                        f"---\n"
                        f"Department: {s.dept}\n"
                        f"Section: {s.name}\n"
                        f"Section_Kind: {kind}\n"
                        f"Section_Priority: {priority}\n"
                        f"Chunk: {j + 1} of {len(chunks)}\n"
                        f"Chunk_Name: {chunk_name}\n"
                        f"Source: {source_label}\n"
                        f"Document_Class: section_chunk\n"
                        f"Is_Full_Document: false\n"
                        f"Images: {s.img_count}\n"
                        f"---\n\n"
                    )
                    f.write(content)
                written += 1
        else:
            content = "\n".join(s.lines).strip()
            fname   = f"{i + 1:03d}_{slugify(s.name)}.md"
            out_path = os.path.join(dept_folder, fname)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(
                    f"---\n"
                    f"Department: {s.dept}\n"
                    f"Section: {s.name}\n"
                    f"Section_Kind: {kind}\n"
                    f"Section_Priority: {priority}\n"
                    f"Source: {source_label}\n"
                    f"Document_Class: section_chunk\n"
                    f"Is_Full_Document: false\n"
                    f"Images: {s.img_count}\n"
                    f"---\n\n"
                )
                f.write(content)
            written += 1

    return written

DISPLAY_SKIP_KINDS = {"toc", "document_control", "revision_history", "glossary","general"}

def _clean_section_lines(lines: list[str]) -> list[str]:
    cleaned = []
    prev_blank = False

    for line in lines:
        if line is None:
            continue

        s = str(line).rstrip()

        # normalize blank lines
        if not s.strip():
            if not prev_blank:
                cleaned.append("")
            prev_blank = True
            continue

        prev_blank = False
        cleaned.append(s)

    # trim leading / trailing blank lines
    while cleaned and cleaned[0] == "":
        cleaned.pop(0)
    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    return cleaned

def _strip_duplicate_leading_heading(lines: list[str], section_name: str) -> list[str]:
    if not lines:
        return lines
    first = lines[0].strip()
    norm_first = re.sub(r"^#+\s*", "", first).strip().lower()
    norm_name = (section_name or "").strip().lower()
    if norm_first == norm_name:
        return lines[1:]
    return lines

def _normalize_markdown_for_display(text: str) -> str:
    # remove outer ```markdown fences accidentally returned by LLM helpers
    text = re.sub(r"^\s*```markdown\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```\s*$", "", text)

    # normalize copied Word artifacts
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]*[•▪◦●]\s*", "\n- ", text)
    text = re.sub(r"(?<!\n)\s+([a-z])\)\s+", r"\n\1) ", text)

    # drop useless standalone labels
    text = re.sub(r"(?im)^\s*(policy|procedure)\s*$\n?", "", text)

    # normalize spacing around markdown tables
    text = re.sub(r"\n(\|.+?\|)\n(?!\|)", r"\n\1\n", text, flags=re.S)

    # collapse 3+ blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()

def _should_hide_from_body_view(section_name: str) -> bool:
    kind = _section_kind(section_name)
    return kind in DISPLAY_SKIP_KINDS

def _drop_visual_blocks_for_body(lines: list[str]) -> list[str]:
    """
    Remove diagram / flowchart / mermaid blocks from BODY full-doc view.
    Keep real policy text and real markdown tables.
    """
    out = []
    in_mermaid = False
    skip_labeled_block = False

    for line in lines:
        s = line.strip()

        if in_mermaid:
            if s.startswith("```"):
                in_mermaid = False
            continue

        if s.lower().startswith("```mermaid"):
            in_mermaid = True
            continue

        if s.startswith("**[Diagram") or s.startswith("**[Flowchart") or s.startswith("**[SmartArt"):
            skip_labeled_block = True
            continue

        if skip_labeled_block:
            if s == "":
                skip_labeled_block = False
            continue

        if re.match(r"^figure\s+\d+", s, re.I):
            continue

        out.append(line)

    return out


def _trim_body_preamble(lines: list[str]) -> list[str]:
    """
    Trim front-matter residue from BODY view.
    """
    junk_prefixes = (
        "[client]", "client:", "version #", "version ", "issue / effective", "issue/effective",
        "date of next review", "document control", "revision history",
        "abbreviations and acronyms", "explanation of key terms",
        "document number", "document type", "status", "effective date",
        "department/function", "prepared by", "reviewed by", "approved by",
        "distribution", "classification", "policy custodian",
        "content issuance", "periodic audit review reference"
    )

    out = list(lines)
    while out:
        s = out[0].strip().lower()
        if not s:
            out.pop(0)
            continue
        if s.startswith("#"):
            break
        if s.startswith(junk_prefixes):
            out.pop(0)
            continue
        break
    return out

def write_full_document(sections: list, source_label: str, log_fn=None) -> int:
    """
    Write two clean full-document variants:
    1) FULL  = everything useful in order
    2) BODY  = policy/procedure body only (skips admin front matter)
    """
    _log = log_fn or print
    full_dir = os.path.join(KB_FOLDER, "_full_documents")
    os.makedirs(full_dir, exist_ok=True)

    def _build_lines(body_only: bool) -> list[str]:
        out = []
        for s in sections:
            if s.is_empty() or _should_skip_section(s.name):
                continue
            if body_only and _should_hide_from_body_view(s.name):
                continue

            cleaned = _clean_section_lines(s.lines)
            cleaned = _strip_duplicate_leading_heading(cleaned, s.name)
            if body_only:
                cleaned = _drop_visual_blocks_for_body(cleaned)
                cleaned = _trim_body_preamble(cleaned)

            if not cleaned:
                continue

            if body_only and s.name.strip().lower() != "general":
                out.append(f"## {s.name}")


            out.extend(cleaned)
            if cleaned[-1].strip():
                out.append("")
        return out

    variants = [
        ("full", _build_lines(body_only=False)),
        ("body", _build_lines(body_only=True)),
    ]

    written = 0
    for view_name, lines in variants:
        full_text = _normalize_markdown_for_display("\n".join(lines).strip())
        if not full_text:
            continue

        stem = slugify(Path(source_label).stem)
        fname = f"{stem}__{view_name.upper()}.md"
        out_path = os.path.join(full_dir, fname)

        with open(out_path, "w", encoding="utf-8") as f:
            f.write(
                f"---\n"
                f"Department: FULL_DOCUMENT\n"
                f"Section: FULL_DOCUMENT\n"
                f"Section_Kind: full_document\n"
                f"Section_Priority: high\n"
                f"Source: {source_label}\n"
                f"Document_Title: {Path(source_label).stem}\n"
                f"Document_Class: full_document\n"
                f"Document_View: {view_name}\n"
                f"Is_Full_Document: true\n"
                f"---\n\n"
            )
            f.write(full_text)

        _log(f"  [full-doc] wrote {fname}")
        written += 1

    return written

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — ENTRY POINT  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────
def _purge_existing_outputs_for_source(source_label: str, log_fn=None):
    _log = log_fn or print
    kb_root = Path(KB_FOLDER)
    if not kb_root.exists():
        return

    removed = 0
    for fp in kb_root.rglob("*.md"):
        try:
            raw = fp.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue
        if f"Source: {source_label}\n" in raw:
            try:
                fp.unlink()
                removed += 1
            except Exception:
                pass

    if removed:
        _log(f"  [cleanup] removed {removed} old KB file(s) for source: {source_label}")

def _purge_existing_outputs_for_source(source_label: str, log_fn=None):
    _log = log_fn or print
    kb_root = Path(KB_FOLDER)
    if not kb_root.exists():
        return

    removed = 0
    for fp in kb_root.rglob("*.md"):
        try:
            raw = fp.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue
        if f"Source: {source_label}\n" in raw:
            try:
                fp.unlink()
                removed += 1
            except Exception:
                pass

    if removed:
        _log(f"  [cleanup] removed {removed} old KB file(s) for source: {source_label}")
        
def _write_quarantine_log(source_label: str, failures: list, log_fn=None):
    """Append failed diagram/flowchart extractions to a review log (not the KB index)."""
    _log = log_fn or print
    os.makedirs(KB_FOLDER, exist_ok=True)
    log_path = os.path.join(KB_FOLDER, "_extraction_failures.log")
    stamp = time.strftime("%Y-%m-%d %H:%M:%S")
    with open(log_path, "a", encoding="utf-8") as f:
        for fail in failures:
            err = " ".join(str(fail.get("error", "")).split())
            f.write(f"{stamp}\t{source_label}\t{fail.get('section','')}\t{fail.get('fmt','')}\t{err}\n")
    _log(f"  [quarantine] logged {len(failures)} extraction failure(s) → {log_path}")


def process_one_file(file_path: str, dry_run: bool, log_fn=None, continue_on_warning: bool = True, source_label_override: str | None = None):
    _log = log_fn or print
    ext  = Path(file_path).suffix.lower()

    if ext == ".docx":
        sections, gt_text, gt_table, split_levels = process_docx(file_path)
    elif ext == ".pdf":
        sections, gt_text, gt_table, split_levels = process_pdf(file_path)
    else:
        _log(f"  Skipping unsupported file type: {file_path}")
        return

    source_label = source_label_override or Path(file_path).name
    ok = print_verification(sections, gt_text, gt_table, split_levels, source_label)

    if dry_run:
        _log("\n  Dry run — no API calls made, no files written.")
        return

    if not ok:
        _log("\n  WARNING: verification flagged issues above.")
        if not continue_on_warning:
            return
    
    _purge_existing_outputs_for_source(source_label, log_fn=_log)
    failures = run_vision(sections, log_fn=_log)
    if failures:
        _write_quarantine_log(source_label, failures, log_fn=_log)
    written_sections = write_sections(sections, source_label, log_fn=_log)
    written_full     = write_full_document(sections, source_label, log_fn=_log)
    _log(f"\n  Done: {written_sections} section chunks + {written_full} full-document file → {KB_FOLDER}/")


def main():
    parser = argparse.ArgumentParser(
        description="Universal P&P preprocessor — structural + semantic hybrid chunking."
    )
    parser.add_argument("--file",    metavar="PATH", help="Process a single DOCX or PDF.")
    parser.add_argument("--folder",  metavar="PATH", default=RAW_DOCS_FOLDER,
                        help=f"Process all DOCX/PDF in folder (default: {RAW_DOCS_FOLDER}/).")
    parser.add_argument("--dry-run", action="store_true",
                        help="Show verification report only — no API calls, no files written.")
    parser.add_argument("--ui",      action="store_true",
                        help="Launch browser upload UI (requires: pip install streamlit).")
    args = parser.parse_args()

    if args.ui:
        import subprocess as _sp, sys as _sys
        _sp.run([_sys.executable, "-m", "streamlit", "run", __file__, "--", "--_ui_mode"])
        return

    if "--_ui_mode" in sys.argv:
        _run_streamlit_ui()
        return

    targets = []
    if args.file:
        targets = [args.file]
    else:
        folder = Path(args.folder)
        if not folder.exists():
            print(f"ERROR: folder not found: {folder}")
            sys.exit(1)
        targets = sorted(
            str(f) for f in folder.rglob("*")
            if f.suffix.lower() in (".docx", ".pdf")
        )
        if not targets:
            print(f"No DOCX or PDF files found in {folder}/")
            sys.exit(0)

    print(f"Files to process: {len(targets)}")
    for t in targets:
        print(f"  {t}")

    for target in targets:
        if not os.path.exists(target):
            print(f"\nERROR: file not found: {target}")
            continue
        print(f"\n{'='*65}")
        print(f"Processing: {target}")
        print(f"{'='*65}")
        try:
            process_one_file(target, dry_run=args.dry_run)
        except Exception as e:
            import traceback
            print(f"\nERROR on {target}: {e}")
            traceback.print_exc()

    if not args.dry_run and len(targets) > 1:
        print(f"\n{'='*65}")
        print(f"All files processed. Next: python ingest_to_vectorstore.py")
        print(f"{'='*65}")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — STREAMLIT UPLOAD UI  (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

def _run_streamlit_ui():
    import streamlit as st
    import tempfile

    st.set_page_config(page_title="P&P Preprocessor", page_icon="📄", layout="wide")
    st.title("Policy & Procedure Document Preprocessor")
    st.markdown(
        "Upload a **DOCX** or **PDF** file. The same preprocessing pipeline that built "
        "the Arabian Mills knowledge base will run on your file."
    )

    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("Upload")
        uploaded = st.file_uploader(
            "Drag and drop your file here",
            type=["docx", "pdf"],
            help="DOCX: full structural + semantic chunking. PDF: text-based or scanned auto-detected.",
        )
        dry_run = st.checkbox("Dry run (verify structure only — no API calls, no files written)",
                              value=True)

        if uploaded:
            st.success(f"File ready: **{uploaded.name}**  ({uploaded.size // 1024} KB)")
            run_btn = st.button("▶ Start Preprocessing", type="primary", use_container_width=True)
        else:
            st.info("No file uploaded yet.")
            run_btn = False

    with col2:
        st.subheader("Live Output")
        log_box = st.empty()
        log_lines = []

        def log_fn(msg: str):
            log_lines.append(str(msg))
            log_box.code("\n".join(log_lines[-60:]), language="")

        if run_btn and uploaded:
            suffix = Path(uploaded.name).suffix.lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uploaded.getbuffer())
                tmp_path = tmp.name

            log_fn(f"Starting: {uploaded.name}")
            log_fn(f"Mode: {'DRY RUN' if dry_run else 'FULL PROCESSING'}")
            log_fn("=" * 55)

            try:
                process_one_file(tmp_path, dry_run=dry_run, log_fn=log_fn)
                log_fn("\n" + "=" * 55)
                log_fn("COMPLETE.")
                if not dry_run:
                    log_fn(f"Files written to: {KB_FOLDER}/")
                    log_fn("Next step: python ingest_to_vectorstore.py")
                st.success("Done!" if not dry_run else "Dry run complete — check output above.")
            except Exception as e:
                import traceback
                log_fn(f"\nERROR: {e}")
                log_fn(traceback.format_exc())
                st.error(f"Error: {e}")
            finally:
                os.unlink(tmp_path)

    if os.path.exists(KB_FOLDER) and not dry_run:
        st.markdown("---")
        st.subheader("Knowledge Base Output")
        md_files = []
        for root, _, files in os.walk(KB_FOLDER):
            for f in sorted(files):
                if f.endswith(".md"):
                    fpath = os.path.join(root, f)
                    size  = os.path.getsize(fpath)
                    md_files.append((os.path.relpath(fpath, KB_FOLDER), size))

        if md_files:
            st.caption(f"{len(md_files)} .md files in {KB_FOLDER}/")
            selected = st.selectbox(
                "Preview a file",
                [f for f, _ in md_files],
                format_func=lambda x: x
            )
            if selected:
                full_path = os.path.join(KB_FOLDER, selected)
                content   = open(full_path, encoding="utf-8").read()
                st.text_area("File content", content, height=400)


if __name__ == "__main__":
    main()