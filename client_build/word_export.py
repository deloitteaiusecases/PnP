"""
word_export.py  —  Requirement 7
=================================
Converts generated policy/procedure text into a professional Word (.docx) document.
Returns raw bytes that Streamlit can serve via st.download_button.

Usage:
    from word_export import generate_word_doc
    docx_bytes = generate_word_doc(policy_text, sources, department="IT", doc_number="POL-001")
    st.download_button("Download .docx", docx_bytes, file_name="policy.docx")
"""

import io, re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _ensure_graphviz_on_path() -> bool:
    """
    Make the `dot` executable findable even when the launching shell didn't have
    Graphviz on PATH (common on Windows when the terminal predates the install).
    Checks PATH first, then well-known install dirs, prepending the first hit.
    """
    import shutil, os
    if shutil.which("dot"):
        return True
    candidates = [
        r"C:\Program Files\Graphviz\bin",
        r"C:\Program Files (x86)\Graphviz\bin",
        os.path.join(os.environ.get("LOCALAPPDATA", ""), r"Programs\Graphviz\bin"),
        "/usr/bin", "/usr/local/bin",  # Linux (Streamlit Cloud installs via packages.txt)
    ]
    for c in candidates:
        if c and (os.path.isfile(os.path.join(c, "dot.exe")) or os.path.isfile(os.path.join(c, "dot"))):
            os.environ["PATH"] = c + os.pathsep + os.environ.get("PATH", "")
            return True
    return bool(shutil.which("dot"))


# High render resolution so the embedded flowchart stays crisp in Word (the default
# ~96 DPI looks pixelated once stretched across the page). 200 DPI is sharp without
# producing an unreasonably large PNG.
_FLOWCHART_DPI = 200


def _inject_dpi(dot_code: str, dpi: int = _FLOWCHART_DPI) -> str:
    """Insert a graph-level dpi attribute after the first '{' so the raster renders
    at high resolution. Leaves the DOT unchanged if it already sets dpi."""
    if re.search(r"\bdpi\s*=", dot_code):
        return dot_code
    return re.sub(r"\{", f'{{ graph [dpi={dpi}]; ', dot_code, count=1)


def _render_dot_to_png(dot_code: str) -> bytes:
    """
    Render a Graphviz DOT string to high-resolution PNG bytes for embedding as a
    flowchart image. Returns b"" if the dot binary is unavailable or the source is
    invalid, so the export never crashes on a malformed/absent diagram.
    """
    if not dot_code.strip():
        return b""
    try:
        _ensure_graphviz_on_path()
        import graphviz
        src = graphviz.Source(_inject_dpi(dot_code), format="png", engine="dot")
        return src.pipe()
    except Exception as e:
        print(f"[word_export] flowchart render skipped: {e}")
        return b""


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, size=9, color=None, align=None):
    p = cell.paragraphs[0]
    p.clear()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def _section_heading(doc, number: str, title: str):
    """ISO-style numbered bold section heading with navy underline rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{number}  ")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    r2 = p.add_run(title.upper())
    r2.bold = True; r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "1F3864")
    pBdr.append(bot); pPr.append(pBdr)


def _set_header_footer(doc, doc_number, title, version):
    for section in doc.sections:
        hdr = section.header
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            p.clear()
        ht = hdr.add_table(1, 3, width=Inches(6.5))
        ht.style = "Table Grid"
        for col in range(3):
            _set_cell_bg(ht.cell(0, col), "1F3864")
        _cell_text(ht.cell(0, 0), "POLICY & PROCEDURE", bold=True, size=8, color="FFFFFF")
        _cell_text(ht.cell(0, 1), (title or "")[:50], size=8, color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(ht.cell(0, 2), f"Doc: {doc_number}  |  Ver: {version}",
                   size=8, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.RIGHT)

        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        r  = fp.add_run(f"CONFIDENTIAL — {doc_number} — Printed copies uncontrolled. "
                        f"Verify revision before use. | Page ")
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for tag, text in [("begin", "PAGE"), ("end", "")]:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), tag)
            it = OxmlElement("w:instrText"); it.text = text if tag == "begin" else ""
            r2 = fp.add_run()
            r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            r2._r.append(fc)
            if tag == "begin":
                r2._r.append(it)

def _infer_doc_type(title: str, policy_text: str = "") -> str:
    title_l = (title or "").lower()
    if "policy" in title_l:
        return "POLICY"
    if "procedure" in title_l:
        return "PROCEDURE"
    if "standard" in title_l:
        return "STANDARD"
    if "sop" in title_l:
        return "SOP"

    combined = f"{title or ''}\n{policy_text or ''}".lower()
    if "policy" in combined:
        return "POLICY"
    if "procedure" in combined:
        return "PROCEDURE"
    if "standard" in combined:
        return "STANDARD"
    if "sop" in combined:
        return "SOP"
    return "POLICY"

def _doc_control_block(doc, doc_number, title, version, department, owner, approver,
                       effective_date, classification, document_type="POLICY"):
    banner = doc.add_table(1, 1)
    banner.style = "Table Grid"
    bc = banner.cell(0, 0)
    _set_cell_bg(bc, "1F3864")
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = bp.add_run(f"{document_type} DOCUMENT\n")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    fallback_title = f"Untitled {document_type.title()}"
    r2 = bp.add_run(title or fallback_title)
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    bp.paragraph_format.space_before = Pt(6)
    bp.paragraph_format.space_after  = Pt(6)
    doc.add_paragraph()

    fields = [
        ("Document Number", doc_number), ("Version", version),
        ("Effective Date", effective_date), ("Department", department or "—"),
        ("Document Owner", owner), ("Approved By", approver),
        ("Classification", classification),
        ("Next Review Date", "12 months from effective date"),
    ]
    ctrl = doc.add_table(len(fields), 2)
    ctrl.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        lc = ctrl.cell(i, 0)
        vc = ctrl.cell(i, 1)
        _set_cell_bg(lc, "D6DCE4")
        _cell_text(lc, label, bold=True, size=9)
        _cell_text(vc, value, size=9)
    doc.add_paragraph()


def _embed_flowchart(doc, dot_code: str):
    """Render DOT to PNG and embed it centered, with a caption. No-op on failure.

    The image is scaled to its TRUE aspect ratio to fit within the printable page
    box (max width and max height), so a tall/narrow flowchart isn't forced to 6"
    wide and blown up off the page. Width is only constrained when the natural size
    (pixels ÷ DPI) exceeds the box."""
    png = _render_dot_to_png(dot_code)
    if not png:
        return
    try:
        from PIL import Image
        with Image.open(io.BytesIO(png)) as im:
            px_w, px_h = im.size
            dpi = (im.info.get("dpi") or (_FLOWCHART_DPI, _FLOWCHART_DPI))[0] or _FLOWCHART_DPI
            nat_w, nat_h = px_w / dpi, px_h / dpi
            aspect = nat_w / nat_h if nat_h else 1.0

            # A horizontal swimlane is wide. Rather than switch the page to landscape
            # (which breaks the document's portrait layout), ROTATE the image 90° so it
            # runs down the portrait page — the reader turns the page sideways to read
            # it. This keeps one consistent page layout and still renders it large.
            rotate = aspect > 1.4
            if rotate:
                rotated = im.rotate(90, expand=True)        # 90° CCW; turn page clockwise to read
                buf = io.BytesIO()
                rotated.save(buf, format="PNG", dpi=(dpi, dpi))
                buf.seek(0)
                img_stream = buf
                nat_w, nat_h = nat_h, nat_w                 # dimensions swap after rotation
            else:
                img_stream = io.BytesIO(png)

        # Fit within the portrait printable area, preserving aspect; never upscale.
        MAX_W, MAX_H = 6.3, 9.0
        scale = min(MAX_W / nat_w, MAX_H / nat_h, 1.0)
        disp_w, disp_h = nat_w * scale, nat_h * scale

        doc.add_picture(img_stream, width=Inches(disp_w), height=Inches(disp_h))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = ("Figure: Process Flowchart (swimlane by responsible role — rotated; view sideways)"
                 if rotate else "Figure: Process Flowchart")
        cr = cap.add_run(label)
        cr.italic = True; cr.font.size = Pt(8)
        cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    except Exception as e:
        print(f"[word_export] flowchart embed skipped: {e}")


def _parse_body(doc, policy_text: str, section_counter: list):
    lines = policy_text.split("\n")
    in_table, table_rows = False, []
    in_fence, fence_lang, fence_lines = False, "", []

    def flush_table():
        nonlocal in_table, table_rows
        data = [r for r in table_rows if not re.match(r"^\s*[\|\-:]+\s*$", r)]
        if not data:
            in_table = False; table_rows = []; return
        cols = len(data[0].strip("|").split("|"))
        t    = doc.add_table(len(data), cols)
        t.style = "Table Grid"
        for ri, row_text in enumerate(data):
            cells = [c.strip() for c in row_text.strip("|").split("|")]
            for ci in range(cols):
                txt  = cells[ci] if ci < len(cells) else ""
                cell = t.cell(ri, ci)
                if ri == 0:
                    _set_cell_bg(cell, "D6DCE4")
                    _cell_text(cell, txt, bold=True, size=9)
                else:
                    bg = "EBF0F7" if ri % 2 == 0 else "FFFFFF"
                    _set_cell_bg(cell, bg)
                    _cell_text(cell, txt, size=9)
        doc.add_paragraph()
        in_table = False; table_rows = []

    for line in lines:
        stripped = line.strip()

        # Fenced code blocks (```dot / ```graphviz / ```mermaid / ```). Capture the
        # inner lines instead of leaking them as paragraphs; render DOT to an image.
        if stripped.startswith("```"):
            if not in_fence:
                if in_table:
                    flush_table()
                in_fence = True
                fence_lang = stripped[3:].strip().lower()
                fence_lines = []
            else:
                if fence_lang in ("dot", "graphviz"):
                    _embed_flowchart(doc, "\n".join(fence_lines))
                # mermaid / other fences: drop quietly (diagrams use DOT here)
                in_fence, fence_lang, fence_lines = False, "", []
            continue
        if in_fence:
            fence_lines.append(line)
            continue

        if stripped.startswith("|"):
            in_table = True; table_rows.append(stripped); continue
        elif in_table:
            flush_table()

        if not stripped:
            continue
        if stripped.startswith("# "):
            section_counter[0] += 1
            _section_heading(doc, str(section_counter[0]), stripped[2:])
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(stripped[3:]); r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        elif stripped.startswith("### "):
            p = doc.add_paragraph(); r = p.add_run(stripped[4:])
            r.bold = True; r.font.size = Pt(10)
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(re.sub(r"^\d+\.\s*", "", stripped)); r.font.size = Pt(10)
        elif stripped.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            content = stripped[2:]
            parts = re.split(r"(\*\*[^*]+\*\*)", content)
            p.clear()
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(10)
                else:
                    r = p.add_run(part); r.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(10)
                else:
                    r = p.add_run(part); r.font.size = Pt(10)

    if in_table:
        flush_table()


def _revision_history(doc, version, effective_date, department):
    headers = ["Version", "Date", "Description", "Prepared By", "Approved By"]
    t = doc.add_table(2, 5); t.style = "Table Grid"
    for ci, h in enumerate(headers):
        c = t.cell(0, ci); _set_cell_bg(c, "1F3864")
        _cell_text(c, h, bold=True, size=9, color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for ci, v in enumerate([version, effective_date, "Initial release", department or "—", "—"]):
        _cell_text(t.cell(1, ci), v, size=9)
    doc.add_paragraph()


def _approval_block(doc, owner, approver):
    headers = ["Role", "Name", "Signature", "Date"]
    rows    = [["Prepared By", owner, "", ""], ["Reviewed By", "—", "", ""],
               ["Approved By", approver, "", ""]]
    t = doc.add_table(1 + len(rows), 4); t.style = "Table Grid"
    for ci, h in enumerate(headers):
        c = t.cell(0, ci); _set_cell_bg(c, "1F3864")
        _cell_text(c, h, bold=True, size=9, color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.cell(ri + 1, ci)
            _set_cell_bg(cell, "F2F2F2" if ci == 2 else "FFFFFF")
            _cell_text(cell, val, size=9)
    doc.add_paragraph()


# ── Public function ───────────────────────────────────────────────────────────

def generate_word_doc(
    policy_text: str,
    sources: list,
    department: str = "",
    doc_number: str  = "POL-001",
    owner: str       = "Policy Owner",
    approver: str    = "Compliance Head",
    version: str     = "1.0",
    classification: str = "Confidential – Internal Use Only",
    document_type: str = "",
) -> bytes:
    """
    Convert generated policy text to a professional ISO-style Word document.
    Returns raw bytes — pass directly to st.download_button.
    """
    doc = Document()

    for section in doc.sections:
        section.page_height   = Cm(29.7)
        section.page_width    = Cm(21.0)
        section.top_margin    = Inches(1.3)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    title_match = re.search(r"^#\s+(.+)$", policy_text, re.MULTILINE)
    doc_title   = title_match.group(1).strip() if title_match else "Policy & Procedure Document"
    final_document_type = (document_type or _infer_doc_type(doc_title, policy_text)).upper()
    eff_date    = date.today().strftime("%d %B %Y")

    _set_header_footer(doc, doc_number, doc_title, version)

    _doc_control_block(
        doc,
        doc_number=doc_number,
        title=doc_title,
        version=version,
        department=department or "—",
        owner=owner,
        approver=approver,
        effective_date=eff_date,
        classification=classification,
        document_type=final_document_type,
    )


    section_counter = [0]
    _parse_body(doc, policy_text, section_counter)

    section_counter[0] += 1
    _section_heading(doc, str(section_counter[0]), "Revision History")
    _revision_history(doc, version, eff_date, department or owner)

    section_counter[0] += 1
    _section_heading(doc, str(section_counter[0]), "Approval & Authorization")
    _approval_block(doc, owner, approver)

    disc = doc.add_paragraph()
    disc.paragraph_format.space_before = Pt(12)
    dr = disc.add_run(
        "DISCLAIMER: Generated by AI-assisted policy tool. "
        "Must be reviewed and approved by authorized personnel before official use."
    )
    dr.italic = True; dr.font.size = Pt(8)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if sources:
        section_counter[0] += 1
        _section_heading(doc, str(section_counter[0]), "Sources Referenced")
        for s in sources:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(s)
            r.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
